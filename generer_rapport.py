# generer_rapport.py - Generation du rapport Word a partir des resultats reels
# Auteur : Papa Malick NDIAYE | Master DSGL, UADB
#
# Le rapport n'est pas ecrit a la main : il est reconstruit depuis
# metrics/results.json et metrics/histories.json. Les chiffres du document
# ne peuvent donc jamais diverger de ceux produits par la pipeline.
#
# Les cles JSON et les identifiants restent en ASCII. Seul le texte
# redige dans le document porte les accents.
#
# Usage :
#   python main.py --split patient      (produit les metriques)
#   python generer_rapport.py

import os
import sys
import json

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "src"))

SORTIE      = "Rapport_Malaria_Detection_NDIAYE.docx"
METRICS_DIR = "metrics"
ASSETS_DIR  = "assets"
DATA_DIR    = "data/cell_images"

DEMO_URL   = "https://deep-learning-malaria-detection-znf2.onrender.com/"
DEPOT_URL  = "https://github.com/pa-malick/MALARIA_DETECTION"

NOIR    = RGBColor(0, 0, 0)
POLICE  = "Calibri"
INSEC   = " "   # espace insecable, separateur de milliers en francais


# ----------------------------------------------------------------------
# Mise en forme
# ----------------------------------------------------------------------

def preparer_document() -> Document:
    """Document sobre : texte noir, une seule police, aucune couleur."""
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = POLICE
    normal.font.size = Pt(11)
    normal.font.color.rgb = NOIR
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # Les styles de titre de Word sont bleus par defaut. On les repasse en noir.
    for niveau, taille in [(1, 15), (2, 12)]:
        style = doc.styles[f"Heading {niveau}"]
        style.font.name = POLICE
        style.font.size = Pt(taille)
        style.font.bold = True
        style.font.color.rgb = NOIR
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.2)
        section.left_margin = section.right_margin = Cm(2.2)

    return doc


def titre(doc, texte: str, niveau: int = 1) -> None:
    doc.add_heading(texte, level=niveau)


def para(doc, texte: str, gras: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(texte)
    r.bold = gras
    r.font.color.rgb = NOIR


def puces(doc, elements: list) -> None:
    for e in elements:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(e)
        r.font.color.rgb = NOIR


def tableau(doc, entetes: list, lignes: list) -> None:
    """Tableau simple, bordures fines, en-tete en gras, aucune trame."""
    t = doc.add_table(rows=1, cols=len(entetes))
    t.style = "Table Grid"

    for cellule, texte in zip(t.rows[0].cells, entetes):
        r = cellule.paragraphs[0].add_run(str(texte))
        r.bold = True
        r.font.color.rgb = NOIR
        r.font.size = Pt(10)

    for ligne in lignes:
        cells = t.add_row().cells
        for cellule, valeur in zip(cells, ligne):
            r = cellule.paragraphs[0].add_run(str(valeur))
            r.font.color.rgb = NOIR
            r.font.size = Pt(10)

    doc.add_paragraph()


def image(doc, nom_fichier: str, legende: str, largeur_cm: float = 14,
          dossier: str = METRICS_DIR) -> None:
    chemin = os.path.join(dossier, nom_fichier)
    if not os.path.exists(chemin):
        return
    doc.add_picture(chemin, width=Cm(largeur_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(legende)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = NOIR


def pct(valeur: float) -> str:
    """Pourcentage au format francais : virgule decimale."""
    return f"{valeur * 100:.2f}".replace(".", ",") + f"{INSEC}%"


def dec(valeur: float, chiffres: int = 1) -> str:
    """Nombre decimal au format francais."""
    return f"{valeur:.{chiffres}f}".replace(".", ",")


def nombre(n) -> str:
    """Entier avec espace insecable comme separateur de milliers."""
    return f"{int(n):,}".replace(",", INSEC)


# ----------------------------------------------------------------------
# Donnees
# ----------------------------------------------------------------------

def charger_resultats() -> dict:
    chemin = os.path.join(METRICS_DIR, "results.json")
    if not os.path.exists(chemin):
        print(f"[ERREUR] {chemin} introuvable.")
        print("Lancez d'abord : python main.py --split patient")
        sys.exit(1)
    with open(chemin, encoding="utf-8") as f:
        return json.load(f)


def charger_historiques() -> dict:
    chemin = os.path.join(METRICS_DIR, "histories.json")
    if not os.path.exists(chemin):
        return {}
    with open(chemin, encoding="utf-8") as f:
        return json.load(f)


def composition_split() -> dict:
    """
    Recalcule la composition du decoupage par patient a partir des noms de
    fichiers. Aucune image n'est ouverte, l'operation est quasi instantanee.
    Retourne un dict vide si le dataset n'est pas present.
    """
    if not os.path.isdir(DATA_DIR):
        return {}
    try:
        from data_loader import charger_chemins
        from preprocessing import split_donnees, extraire_patient
    except ImportError:
        return {}

    import contextlib
    import io as _io
    with contextlib.redirect_stdout(_io.StringIO()):
        chemins, labels = charger_chemins(DATA_DIR)
        tr, va, te = split_donnees(chemins, labels, seed=42, par_patient=True)

    infos = {}
    for nom, df in [("train", tr), ("validation", va), ("test", te)]:
        infos[nom] = {
            "images": len(df),
            "patients": len({extraire_patient(f) for f in df["filename"]}),
            "parasitees": (df["class"] == "Parasitized").mean(),
        }
    infos["total_patients"] = len({extraire_patient(c) for c in chemins})
    infos["total_images"]   = len(chemins)
    return infos


# ----------------------------------------------------------------------
# Chapitres
# ----------------------------------------------------------------------

def page_titre(doc) -> None:
    logo = os.path.join(ASSETS_DIR, "logo_uadb.png")
    if os.path.exists(logo):
        doc.add_picture(logo, width=Cm(9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DÉTECTION AUTOMATIQUE DU PALUDISME\nPAR APPRENTISSAGE PROFOND")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = NOIR

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Classification d'images de cellules sanguines\n"
                  "par réseaux de neurones convolutifs")
    r.font.size = Pt(12)
    r.font.color.rgb = NOIR

    for _ in range(6):
        doc.add_paragraph()

    for texte, taille, gras in [
        ("Papa Malick NDIAYE", 13, True),
        ("Master Data Science et Génie Logiciel", 11, False),
        ("Université Alioune Diop de Bambey", 11, False),
        ("Année académique 2024-2025", 11, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(texte)
        r.bold = gras
        r.font.size = Pt(taille)
        r.font.color.rgb = NOIR

    doc.add_paragraph()
    for etiquette, url in [("Démonstration en ligne : ", DEMO_URL),
                           ("Code source : ", DEPOT_URL)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(etiquette + url)
        r.font.size = Pt(9)
        r.font.color.rgb = NOIR

    doc.add_page_break()


def resume(doc, d: dict) -> None:
    titre(doc, "Résumé")

    meilleur = d["meilleur_modele"]
    m = d["resultats"][meilleur]
    v = d.get("validation", {}).get(meilleur, {})

    para(doc,
         "Ce projet classe des images de cellules sanguines en deux catégories : "
         "cellule parasitée par Plasmodium, ou cellule saine. Trois architectures "
         "de réseaux de neurones convolutifs sont comparées sur le jeu de données "
         f"NIH, qui contient {nombre(27558)} images.")

    para(doc,
         f"Le modèle retenu est {meilleur}. Sur le jeu de test, il atteint une "
         f"accuracy de {pct(m['accuracy'])} et une sensibilité de "
         f"{pct(m['sensibilite'])}. La sensibilité est la métrique importante "
         "ici : elle mesure la proportion de cellules malades effectivement "
         "détectées.")

    if v:
        ecart = abs(m["sensibilite"] - v["sensibilite"]) * 100
        para(doc,
             "Ce chiffre doit être lu avec prudence. Sur le jeu de validation, "
             "composé d'autres patients, la sensibilité du même modèle tombe à "
             f"{pct(v['sensibilite'])}. L'écart entre les deux mesures atteint "
             f"{dec(ecart)} points. La performance réelle dépend donc fortement "
             "des patients évalués. Le chapitre 7 revient sur ce point.")

    para(doc,
         "Le système est accompagné d'une interface web, d'une API REST, d'une "
         "image Docker et d'une suite de tests automatisés.")

    para(doc,
         "Mots-clés : apprentissage profond, réseaux convolutifs, paludisme, "
         "imagerie médicale, TensorFlow, Flask, Docker.")


def chapitre_contexte(doc) -> None:
    titre(doc, "1. Contexte et objectifs")

    titre(doc, "1.1 Le problème", 2)
    para(doc,
         "Le paludisme est une maladie parasitaire transmise par le moustique "
         "anophèle. Il reste une cause majeure de mortalité en Afrique "
         "subsaharienne, et le Sénégal figure parmi les pays touchés.")
    para(doc,
         "Le diagnostic de référence se fait au microscope. Un technicien "
         "examine un frottis sanguin et cherche le parasite dans les globules "
         "rouges. Cette méthode est fiable mais lente. Elle demande un opérateur "
         "formé, et sa qualité varie selon l'expérience de celui-ci.")

    titre(doc, "1.2 L'idée", 2)
    para(doc,
         "Un réseau de neurones convolutif sait classer des images. On peut donc "
         "lui apprendre à distinguer une cellule parasitée d'une cellule saine, "
         "à partir d'exemples déjà étiquetés par des experts.")
    para(doc,
         "L'objectif n'est pas de remplacer le microscopiste. Il est de fournir "
         "une aide au tri, qui signale les cellules suspectes.")

    titre(doc, "1.3 Objectifs du projet", 2)
    puces(doc, [
        "Construire une pipeline complète, du chargement des images à l'évaluation.",
        "Comparer trois architectures convolutives de complexité croissante.",
        "Mesurer la performance avec un protocole d'évaluation honnête.",
        "Exposer le modèle via une API et une interface web.",
        "Rendre le tout reproductible avec Docker et des tests automatisés.",
    ])


def chapitre_donnees(doc, split: dict) -> None:
    titre(doc, "2. Les données")

    titre(doc, "2.1 Le jeu NIH", 2)
    para(doc,
         "Le jeu de données provient de la National Library of Medicine des "
         f"États-Unis. Il contient {nombre(27558)} images de cellules sanguines "
         "isolées, extraites de frottis minces colorés au Giemsa.")
    para(doc,
         f"Il est parfaitement équilibré : {nombre(13779)} cellules parasitées "
         f"et {nombre(13779)} cellules saines. Cet équilibre est confortable, "
         "car il évite d'avoir à corriger un déséquilibre de classes.")

    tableau(doc,
            ["Caractéristique", "Valeur"],
            [["Nombre total d'images", nombre(27558)],
             ["Cellules parasitées", nombre(13779)],
             ["Cellules saines", nombre(13779)],
             ["Format", "PNG, couleur"],
             ["Taille d'entrée du réseau", "64 x 64 pixels"]])

    image(doc, "exemples_images.png",
          "Figure 1. Exemples d'images du jeu de données.")

    titre(doc, "2.2 Une information cachée dans les noms de fichiers", 2)
    para(doc,
         "Les noms de fichiers suivent un format précis, par exemple "
         "C116P77ThinF_IMG_20150930_171219_cell_110.png. Le préfixe C116P77 "
         "identifie le frottis, donc le patient.")
    para(doc,
         "Cette information est essentielle. Le jeu ne contient pas "
         f"{nombre(27558)} prélèvements indépendants, mais les cellules de 200 "
         "patients seulement. Un même patient fournit entre 65 et 702 cellules.")

    if split:
        para(doc,
             f"Le projet compte {split['total_patients']} patients pour "
             f"{nombre(split['total_images'])} images.")


def chapitre_protocole(doc, split: dict) -> None:
    titre(doc, "3. Le protocole d'évaluation")

    titre(doc, "3.1 Pourquoi le découpage habituel est trompeur", 2)
    para(doc,
         "La pratique courante consiste à mélanger toutes les images, puis à en "
         "tirer 70 % pour l'entraînement, 15 % pour la validation et 15 % pour "
         "le test.")
    para(doc,
         "Appliqué ici, ce découpage produit un problème grave. Comme un patient "
         "fournit des dizaines de cellules, ses cellules se retrouvent des deux "
         "côtés du découpage. La mesure effectuée sur ce projet est sans appel : "
         "100 % des images de test proviennent de patients également présents à "
         "l'entraînement, et les 200 patients sont partagés.")
    para(doc,
         "Le modèle peut alors obtenir un bon score en reconnaissant un patient "
         "déjà vu, et non en reconnaissant le parasite. Le score mesuré ne dit "
         "plus ce qu'on croit qu'il dit.")

    titre(doc, "3.2 Le découpage par patient", 2)
    para(doc,
         "La correction consiste à découper par patient et non par image. Chaque "
         "patient appartient à un seul ensemble. Le modèle est ainsi évalué sur "
         "des patients qu'il n'a jamais vus, ce qui correspond à son usage réel.")

    if split:
        tableau(doc,
                ["Ensemble", "Patients", "Images", "Part de cellules parasitées"],
                [[nom.capitalize(),
                  split[nom]["patients"],
                  nombre(split[nom]["images"]),
                  dec(split[nom]["parasitees"] * 100) + f"{INSEC}%"]
                 for nom in ["train", "validation", "test"]])

    para(doc,
         "C'est ce protocole qui est utilisé pour tous les résultats présentés "
         "dans ce rapport.")

    titre(doc, "3.3 Les métriques retenues", 2)
    para(doc,
         "La classe positive est la cellule parasitée. Ce choix a une "
         "conséquence directe sur le sens des métriques.")

    tableau(doc,
            ["Métrique", "Ce qu'elle mesure", "Pourquoi elle compte"],
            [["Sensibilité",
              "Part des cellules malades détectées",
              "Un manque signifie un malade non repéré"],
             ["Spécificité",
              "Part des cellules saines bien classées",
              "Un manque signifie une fausse alerte"],
             ["Précision",
              "Part des alertes qui sont justes",
              "Mesure la fiabilité d'une alerte"],
             ["Accuracy",
              "Part totale de bonnes réponses",
              "Utile mais insuffisante seule"]])

    para(doc,
         "En dépistage, la sensibilité prime. Une fausse alerte entraîne un "
         "examen supplémentaire. Un malade non détecté rentre chez lui sans "
         "traitement.")


def chapitre_architectures(doc) -> None:
    titre(doc, "4. Les architectures comparées")

    titre(doc, "4.1 Principe d'un réseau convolutif", 2)
    para(doc,
         "Un réseau convolutif analyse une image par petites zones. Un filtre de "
         "3 x 3 pixels balaie l'image et réagit à un motif précis : un bord, une "
         "texture, une tache sombre.")
    para(doc,
         "Les premières couches détectent des motifs simples. Les couches "
         "suivantes combinent ces motifs pour former des formes plus complexes. "
         "C'est ce qui permet au réseau de reconnaître le parasite, qui apparaît "
         "comme une tache colorée dans le globule rouge.")
    para(doc,
         "Entre deux blocs de convolution, une couche de sous-échantillonnage "
         "réduit la taille de l'image. Le réseau perd du détail mais gagne en "
         "capacité à généraliser.")

    titre(doc, "4.2 Les trois modèles", 2)
    tableau(doc,
            ["Modèle", "Blocs de convolution", "Particularité"],
            [["CNN_Simple", "2", "Référence de base, peu de paramètres"],
             ["CNN_Deep", "3", "Plus profond, plus de capacité"],
             ["CNN_BN", "3", "Ajout de Batch Normalization"]])

    para(doc,
         "Les trois partagent la même entrée de 64 x 64 pixels en couleur et la "
         "même sortie : une valeur entre 0 et 1, interprétée comme la "
         "probabilité que la cellule soit saine.")

    titre(doc, "4.3 Lutte contre le surapprentissage", 2)
    para(doc,
         "Le surapprentissage survient quand le modèle mémorise les images "
         "d'entraînement au lieu d'apprendre à généraliser. Quatre mécanismes "
         "sont utilisés pour le limiter.")
    puces(doc, [
        "Dropout : désactive au hasard une partie des neurones à chaque passage.",
        "Batch Normalization : stabilise les valeurs circulant entre les couches.",
        "Augmentation des données : rotations, décalages, symétries et zooms.",
        "Arrêt anticipé : stoppe l'entraînement dès que la validation se dégrade.",
    ])
    para(doc,
         "L'augmentation est appliquée uniquement aux données d'entraînement. "
         "Les jeux de validation et de test restent intacts, sans quoi la mesure "
         "serait faussée.")


def chapitre_entrainement(doc, historiques: dict) -> None:
    titre(doc, "5. L'entraînement")

    para(doc,
         "Les trois modèles sont entraînés dans les mêmes conditions, avec la "
         "même graine aléatoire, ce qui rend la comparaison équitable et le run "
         "reproductible.")

    tableau(doc,
            ["Paramètre", "Valeur"],
            [["Optimiseur", "Adam"],
             ["Fonction de coût", "Binary crossentropy"],
             ["Taille de lot", "32"],
             ["Époques maximum", "20"],
             ["Graine aléatoire", "42"]])

    titre(doc, "5.1 Les callbacks", 2)
    tableau(doc,
            ["Callback", "Rôle"],
            [["EarlyStopping",
              "Arrête l'entraînement après 5 époques sans progrès"],
             ["ModelCheckpoint",
              "Enregistre le modèle à son meilleur état"],
             ["ReduceLROnPlateau",
              "Diminue le pas d'apprentissage en cas de stagnation"]])

    para(doc,
         "Les trois callbacks surveillent la même grandeur, la perte de "
         "validation. Ce détail évite une incohérence classique : si l'arrêt "
         "anticipé et la sauvegarde suivaient des critères différents, le "
         "fichier enregistré ne correspondrait pas au modèle évalué.")

    if historiques:
        titre(doc, "5.2 Déroulement observé", 2)
        lignes = []
        for nom, h in historiques.items():
            epoques  = len(h.get("loss", []))
            val_accs = h.get("val_accuracy", [])
            lignes.append([nom, epoques, pct(max(val_accs)) if val_accs else "-"])
        tableau(doc,
                ["Modèle", "Époques effectuées", "Meilleure accuracy en validation"],
                lignes)
        para(doc,
             "Les écarts d'époques entre modèles viennent de l'arrêt anticipé, "
             "qui se déclenche à des moments différents selon la convergence.")

    image(doc, "learning_curves_CNN_Deep.png",
          "Figure 2. Courbes d'apprentissage du modèle CNN_Deep.")


def chapitre_resultats(doc, d: dict) -> None:
    titre(doc, "6. Résultats")

    meilleur = d["meilleur_modele"]

    para(doc,
         "Le modèle est choisi sur le jeu de validation. Le jeu de test ne sert "
         "qu'à la mesure finale. Choisir le modèle sur le test reviendrait à "
         "s'auto-attribuer une bonne note.")

    titre(doc, "6.1 Performances sur le jeu de test", 2)
    tableau(doc,
            ["Modèle", "Accuracy", "Sensibilité", "Spécificité", "Précision", "F1"],
            [[nom + (" (retenu)" if nom == meilleur else ""),
              pct(m["accuracy"]), pct(m["sensibilite"]), pct(m["specificite"]),
              pct(m["precision"]), pct(m["f1_score"])]
             for nom, m in d["resultats"].items()])

    image(doc, "comparaison_modeles.png",
          "Figure 3. Comparaison des trois modèles sur le jeu de test.")

    titre(doc, "6.2 Lecture de la matrice de confusion", 2)
    para(doc,
         "La matrice de confusion détaille les erreurs. La case importante est "
         "celle des cellules parasitées classées comme saines : ce sont les "
         "malades manqués.")

    image(doc, f"cm_{meilleur}.png",
          f"Figure 4. Matrice de confusion du modèle {meilleur}.")

    titre(doc, "6.3 Les trois modèles se valent", 2)
    accuracies = [m["accuracy"] for m in d["resultats"].values()]
    ecart = (max(accuracies) - min(accuracies)) * 100
    para(doc,
         "L'écart d'accuracy entre le meilleur et le moins bon modèle est de "
         f"{dec(ecart)} point. Chaque architecture n'a été entraînée qu'une "
         "fois. Un tel écart ne permet pas de conclure qu'une architecture est "
         "supérieure aux autres.")
    para(doc,
         "La conclusion utile est inverse de celle attendue : la profondeur "
         "supplémentaire n'apporte rien sur ce problème. À performance égale, "
         "le modèle le plus léger est préférable.")


def chapitre_limites(doc, d: dict) -> None:
    titre(doc, "7. Limites")

    para(doc,
         "Cette section est volontairement détaillée. Un résultat dont on ne "
         "connaît pas les limites n'est pas exploitable.")

    meilleur = d["meilleur_modele"]
    m = d["resultats"][meilleur]
    v = d.get("validation", {}).get(meilleur, {})

    titre(doc, "7.1 L'estimation est instable", 2)
    if v:
        tableau(doc,
                ["Mesure", "Jeu de validation", "Jeu de test"],
                [["Accuracy", pct(v["accuracy"]), pct(m["accuracy"])],
                 ["Sensibilité", pct(v["sensibilite"]), pct(m["sensibilite"])]])

        ecart = abs(m["sensibilite"] - v["sensibilite"]) * 100
        bas   = min(v["sensibilite"], m["sensibilite"]) * 100
        haut  = max(v["sensibilite"], m["sensibilite"]) * 100

        para(doc,
             "Il s'agit du même modèle, du même entraînement, du même protocole. "
             f"Seuls les patients changent. L'écart de sensibilité atteint "
             f"{dec(ecart)} points.")
        para(doc,
             "Avec une trentaine de patients par ensemble, un chiffre unique "
             "n'est pas fiable. La formulation honnête est que la sensibilité du "
             f"modèle se situe entre {dec(bas, 0)}{INSEC}% et {dec(haut, 0)}"
             f"{INSEC}% selon les patients.")
        para(doc,
             "Une validation croisée par groupes de patients donnerait une "
             "moyenne et un écart-type. Elle n'a pas été réalisée, son coût "
             "étant d'environ cinq entraînements complets.")

    titre(doc, "7.2 Le découpage par image masquait cette instabilité", 2)
    para(doc,
         "Avec l'ancien découpage, validation et test donnaient des résultats "
         "quasiment identiques. Cette concordance ressemblait à de la "
         "robustesse. Elle n'en était pas : les deux ensembles contenaient les "
         "mêmes patients, ils ne pouvaient que s'accorder. La fuite ne gonflait "
         "pas seulement le score, elle dissimulait la dispersion réelle.")

    titre(doc, "7.3 Autres limites", 2)
    puces(doc, [
        "Le modèle manque une part non négligeable des cellules parasitées, "
        "ce qui est incompatible avec un usage clinique.",
        "La validation se limite au jeu NIH : frottis minces, coloration Giemsa, "
        "images réduites à 64 x 64 pixels. Rien ne garantit la transposition à "
        "d'autres protocoles ou d'autres microscopes.",
        "Chaque architecture n'a été entraînée qu'une seule fois.",
        "Aucune validation clinique n'a été conduite.",
    ])

    titre(doc, "7.4 Avertissement", 2)
    para(doc,
         "Ce travail est un exercice académique. Il ne constitue pas un "
         "dispositif médical et ne remplace pas un diagnostic posé par un "
         "professionnel de santé.", gras=True)


def chapitre_application(doc) -> None:
    titre(doc, "8. Application et déploiement")

    titre(doc, "8.1 L'API", 2)
    tableau(doc,
            ["Route", "Méthode", "Rôle"],
            [["/", "GET", "Interface web"],
             ["/predict", "POST", "Analyse une image et retourne la prédiction"],
             ["/metrics", "GET", "Métriques des modèles au format JSON"],
             ["/health", "GET", "État du service"]])

    titre(doc, "8.2 Traitement d'une image", 2)
    para(doc,
         "L'image envoyée est validée, redimensionnée en 64 x 64 pixels, puis "
         "normalisée. Le modèle retourne une probabilité, convertie en verdict.")
    para(doc,
         "L'image est supprimée du serveur immédiatement après l'analyse. Une "
         "taille maximale est imposée. Ces deux points évitent la saturation du "
         "disque et la conservation involontaire de données personnelles.")

    titre(doc, "8.3 L'interface", 2)
    para(doc,
         "L'utilisateur dépose une image de cellule et obtient un verdict "
         "accompagné de la confiance du modèle. L'avertissement rappelant que "
         "l'outil n'est pas un dispositif médical est affiché en permanence "
         "au-dessus de la zone d'analyse.")

    image(doc, "capture_analyse.png",
          "Figure 5. Analyse d'une cellule parasitée dans l'interface web.",
          dossier=ASSETS_DIR)

    para(doc,
         "L'interface expose également le tableau comparatif des trois modèles, "
         "alimenté par le fichier de métriques. Les colonnes reprennent la "
         "sensibilité et la spécificité, et non des métriques génériques.")

    image(doc, "capture_modeles.png",
          "Figure 6. Comparaison des modèles telle qu'affichée par l'application.",
          dossier=ASSETS_DIR)

    titre(doc, "8.4 Déploiement", 2)
    para(doc,
         "L'application est empaquetée dans une image Docker et servie par "
         "gunicorn. Le jeu de données est exclu de l'image, qui n'a besoin que "
         "du modèle entraîné. Le conteneur ne s'exécute pas en tant que "
         "superutilisateur.")
    para(doc,
         f"L'application est déployée et accessible publiquement à l'adresse "
         f"suivante : {DEMO_URL}")


def chapitre_tests(doc) -> None:
    titre(doc, "9. Tests et qualité")

    para(doc,
         "Le projet est couvert par une suite de tests automatisés, exécutée à "
         "chaque modification par une chaîne d'intégration continue.")

    tableau(doc,
            ["Fichier", "Ce qui est vérifié"],
            [["test_models.py", "Formes de sortie et structure des trois réseaux"],
             ["test_preprocessing.py",
              "Découpages, et surtout l'étanchéité du découpage par patient"],
             ["test_utils.py", "Préparation d'image et format des prédictions"],
             ["test_api.py", "Routes, rejets d'upload et suppression des fichiers"]])

    para(doc,
         "Le test le plus important vérifie qu'aucun patient n'apparaît dans "
         "deux ensembles à la fois. Il empêche la fuite décrite au chapitre 3 "
         "de revenir sans être remarquée.")

    para(doc,
         "Les tests s'exécutent sans le jeu de données ni les modèles entraînés, "
         "ce qui permet de les lancer automatiquement sur un serveur "
         "d'intégration.")


def conclusion(doc, d: dict) -> None:
    titre(doc, "10. Conclusion")

    meilleur = d["meilleur_modele"]
    m = d["resultats"][meilleur]

    para(doc,
         f"Le projet aboutit à un système fonctionnel. Le modèle {meilleur} "
         f"atteint {pct(m['accuracy'])} d'accuracy et {pct(m['sensibilite'])} de "
         "sensibilité sur des patients jamais vus à l'entraînement.")

    para(doc,
         "L'apport principal de ce travail n'est pas le score. Il est la mise en "
         "évidence d'un biais méthodologique : le découpage habituel des données "
         "mélange les patients entre entraînement et test, ce qui produit des "
         "chiffres flatteurs et masque la variabilité réelle entre patients.")

    para(doc,
         "Corriger ce biais a changé les conclusions du projet. Le modèle retenu "
         "n'est plus le même, l'écart entre architectures s'est révélé non "
         "significatif, et la performance s'exprime désormais comme un "
         "intervalle plutôt que comme un chiffre unique.")

    titre(doc, "10.1 Perspectives", 2)
    puces(doc, [
        "Validation croisée par groupes de patients, pour une moyenne et un écart-type.",
        "Ajustement du seuil de décision pour privilégier la sensibilité.",
        "Transfer learning à partir d'un modèle pré-entraîné.",
        "Grad-CAM, pour visualiser les zones observées par le réseau.",
        "Évaluation sur un jeu de données issu d'un autre laboratoire.",
    ])


def references(doc) -> None:
    titre(doc, "Références")
    puces(doc, [
        "National Library of Medicine. Malaria Datasets. "
        "https://ceb.nlm.nih.gov/repositories/malaria-datasets/",
        "Rajaraman S. et collaborateurs. Pre-trained convolutional neural "
        "networks as feature extractors toward improved malaria parasite "
        "detection. PeerJ, 2018.",
        "Organisation mondiale de la santé. World Malaria Report.",
        "Documentation TensorFlow et Keras. https://www.tensorflow.org/",
        "Documentation scikit-learn. https://scikit-learn.org/",
    ])


# ----------------------------------------------------------------------

def main() -> None:
    d           = charger_resultats()
    historiques = charger_historiques()
    split       = composition_split()

    doc = preparer_document()

    page_titre(doc)
    resume(doc, d)
    chapitre_contexte(doc)
    chapitre_donnees(doc, split)
    chapitre_protocole(doc, split)
    chapitre_architectures(doc)
    chapitre_entrainement(doc, historiques)
    chapitre_resultats(doc, d)
    chapitre_limites(doc, d)
    chapitre_application(doc)
    chapitre_tests(doc)
    conclusion(doc, d)
    references(doc)

    doc.save(SORTIE)
    print(f"Rapport genere : {SORTIE}")
    print(f"  modele retenu : {d['meilleur_modele']}")
    print(f"  chapitres : 10, plus resume et references")


if __name__ == "__main__":
    main()
